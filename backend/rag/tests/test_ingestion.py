import os
import json
import pytest
import fitz
import docx
import openpyxl
from unittest.mock import MagicMock, patch
from rag.processing.metadata import MetadataExtractor
from rag.ingestion.registry import LoaderRegistry
from rag.ingestion.text_loader import TextLoader, CSVLoader, JSONLoader
from rag.ingestion.pdf_loader import PDFLoader
from rag.ingestion.docx_loader import DocxLoader
from rag.ingestion.xlsx_loader import XLSXLoader
from rag.pipeline.ingestion_pipeline import IngestionPipeline

class TestMetadataExtractor:
    def test_sha256_computation(self, tmp_path):
        test_file = tmp_path / "test.txt"
        test_file.write_text("Deterministic content for SHA256", encoding="utf-8")
        hash1 = MetadataExtractor.compute_sha256(str(test_file))
        hash2 = MetadataExtractor.compute_sha256(str(test_file))
        assert hash1 == hash2
        assert len(hash1) == 64

    def test_metadata_extraction_heuristics(self, tmp_path):
        test_file = tmp_path / "SOP-042 Valve V-204 Manual.pdf"
        test_file.write_text("Sample", encoding="utf-8")
        
        meta = MetadataExtractor.extract_metadata(str(test_file), test_file.name)
        assert meta.document_id == "SOP-042"
        assert meta.equipment_id == "V-204"
        assert meta.document_type == "SOP"
        assert len(meta.sha256_hash) == 64

    def test_metadata_override(self, tmp_path):
        test_file = tmp_path / "generic_doc.txt"
        test_file.write_text("Sample", encoding="utf-8")
        
        override = {
            "document_id": "CUSTOM-999",
            "equipment_id": "P-500",
            "document_type": "CALIBRATION_CERT",
            "department": "instrumentation",
            "revision": "v4"
        }
        meta = MetadataExtractor.extract_metadata(str(test_file), test_file.name, override)
        assert meta.document_id == "CUSTOM-999"
        assert meta.equipment_id == "P-500"
        assert meta.document_type == "CALIBRATION_CERT"
        assert meta.department == "instrumentation"
        assert meta.revision == "v4"


class TestLoaders:
    def test_loader_registry(self):
        reg = LoaderRegistry()
        assert isinstance(reg.get_loader("doc.pdf"), PDFLoader)
        assert isinstance(reg.get_loader("doc.PDF"), PDFLoader)
        assert isinstance(reg.get_loader("doc.docx"), DocxLoader)
        assert isinstance(reg.get_loader("data.xlsx"), XLSXLoader)
        assert isinstance(reg.get_loader("notes.txt"), TextLoader)
        assert isinstance(reg.get_loader("readme.md"), TextLoader)
        assert isinstance(reg.get_loader("table.csv"), CSVLoader)
        assert isinstance(reg.get_loader("data.json"), JSONLoader)

        with pytest.raises(ValueError):
            reg.get_loader("archive.zip")

    def test_text_loader(self, tmp_path):
        f = tmp_path / "sample.txt"
        f.write_text("Line 1\nLine 2\nLine 3", encoding="utf-8")
        loader = TextLoader()
        pages = loader.load(str(f))
        assert len(pages) == 1
        assert "Line 1" in pages[0].text

    def test_csv_loader(self, tmp_path):
        f = tmp_path / "sample.csv"
        f.write_text("Equip,Status,Temp\nV-204,OK,75\nP-101,WARN,95", encoding="utf-8")
        loader = CSVLoader()
        pages = loader.load(str(f))
        assert len(pages) == 1
        assert "V-204 | OK | 75" in pages[0].text
        assert pages[0].metadata["rows"] == 3

    def test_json_loader(self, tmp_path):
        f = tmp_path / "sample.json"
        data = {"equipment_id": "V-204", "status": "nominal"}
        f.write_text(json.dumps(data), encoding="utf-8")
        loader = JSONLoader()
        pages = loader.load(str(f))
        assert len(pages) == 1
        assert '"equipment_id": "V-204"' in pages[0].text

    def test_pdf_loader(self, tmp_path):
        f = str(tmp_path / "sample.pdf")
        doc = fitz.open()
        p1 = doc.new_page()
        p1.insert_text((50, 72), "PDF Page 1 Content")
        p2 = doc.new_page()
        p2.insert_text((50, 72), "PDF Page 2 Content")
        doc.save(f)
        doc.close()

        loader = PDFLoader()
        pages = loader.load(f)
        assert len(pages) == 2
        assert pages[0].page_number == 1
        assert "PDF Page 1 Content" in pages[0].text
        assert pages[1].page_number == 2
        assert "PDF Page 2 Content" in pages[1].text

    def test_docx_loader(self, tmp_path):
        f = str(tmp_path / "sample.docx")
        doc = docx.Document()
        doc.add_paragraph("Paragraph 1 in Word Doc")
        doc.add_paragraph("Paragraph 2 in Word Doc")
        doc.save(f)

        loader = DocxLoader()
        pages = loader.load(f)
        assert len(pages) == 1
        assert "Paragraph 1 in Word Doc" in pages[0].text
        assert "Paragraph 2 in Word Doc" in pages[0].text

    def test_xlsx_loader(self, tmp_path):
        f = str(tmp_path / "sample.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "SensorReadings"
        ws.append(["Sensor", "Value", "Unit"])
        ws.append(["PT-101", 14.5, "bar"])
        wb.save(f)
        wb.close()

        loader = XLSXLoader()
        pages = loader.load(f)
        assert len(pages) == 1
        assert "Sheet: SensorReadings" in pages[0].text
        assert "PT-101 | 14.5 | bar" in pages[0].text


class TestIngestionPipelineWorkflow:
    def test_ingestion_and_duplicate_skip(self, tmp_path):
        doc_file = str(tmp_path / "SOP-042 Valve Procedure.txt")
        with open(doc_file, "w", encoding="utf-8") as f:
            f.write("1.0 Valve Inspection\nInspect valve V-204 seal seat for wear.")

        with patch("rag.pipeline.ingestion_pipeline.LocalEmbeddingService.get_instance") as mock_embed_inst, \
             patch("rag.pipeline.ingestion_pipeline.QdrantStore") as mock_store_cls:
            
            mock_store = MagicMock()
            mock_store.check_document_hash_exists.return_value = False
            mock_store_cls.return_value = mock_store

            mock_embed = MagicMock()
            mock_embed.embed_batch.return_value = [[0.1] * 384]
            mock_embed_inst.return_value = mock_embed

            pipeline = IngestionPipeline()
            pipeline.store = mock_store
            pipeline.embedder = mock_embed

            # First Ingestion
            res1 = pipeline.ingest_file(doc_file)
            assert res1["status"] == "SUCCESS"
            assert res1["document_id"] == "SOP-042"
            assert res1["chunks_indexed"] > 0
            assert mock_store.upsert_chunks.called

            # Duplicate Ingestion
            mock_store.check_document_hash_exists.return_value = True
            res2 = pipeline.ingest_file(doc_file)
            assert res2["status"] == "SKIPPED"
            assert "Duplicate" in res2["reason"]
